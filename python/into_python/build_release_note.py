#!/usr/local/bin/python3
#-*- coding:utf-8 -*-
#Filename: test_cb.py
#System:   14.04.1-Ubuntu

import os
import docx

release_info_txt = r"./Release_info.txt"
release_note_docx = r"./Release_note.docx"

# Read data from release_info.txt
with open(release_info_txt, "r") as f:
	datalines = f.readlines()
	print ("release info: {}".format(datalines))


docx_file = docx.Document(release_note_docx)

print ("Paragraph number: {}".format(len(docx_file.paragraphs)))
for i in range(len(docx_file.paragraphs)):
	print("line"+str(i)+"："+docx_file.paragraphs[i].text)

# Add text of Kill bugs:
for i in range(len(docx_file.paragraphs)):
	if docx_file.paragraphs[i].text == "Killed Bugs":
		print ("Killed Bugs line number: {}".format(i))

		pos = datalines.index("Killed Bugs:\n")
		for j in range(len(datalines) - pos):
			p = docx_file.paragraphs[i].clear()
			p.text = datalines[pos+j]
			#docx_file.add_paragraph(datalines[pos + j + 1])
			
docx_file.save("./Release Note.docx")




